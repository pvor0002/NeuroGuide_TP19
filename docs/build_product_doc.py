"""Generate NeuroGuide Product Document as Word (.docx)."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

import sys

OUT = sys.argv[1] if len(sys.argv) > 1 else "PRODUCT_DOCUMENT.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("NeuroGuide — Product Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Handover Package 2026 S1 · Team Stellar Minds (TP19)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. Executive overview", 1)
    add_para(
        doc,
        "NeuroGuide is a web-based career support tool for early-career IT professionals "
        "and recent graduates with ADHD (and related attention or information-overload challenges). "
        "It reduces cognitive load when exploring jobs by turning dense job descriptions into "
        "plain-language summaries, estimating role–profile fit with transparent scoring, and guiding "
        "structured interview preparation and “day in the life” exploration.",
    )
    add_para(doc, "Why it exists: ", bold=True)
    p = doc.paragraphs[-1]
    p.add_run(
        "Job postings are written for neurotypical scanning patterns—long blocks, implicit "
        "requirements, and jargon. Candidates with ADHD often experience overload, miss constraints, "
        "or avoid applying despite being qualified. NeuroGuide does not replace employers, clinicians, "
        "or career counsellors; it supports informed self-directed decisions in the application pipeline."
    )
    add_para(doc, "Audience: Primary users are early-career IT job seekers with ADHD; secondary users include anyone who benefits from chunked, accessible job information.")
    add_para(doc, "Why it matters: Neurodivergent talent is under-represented in IT hiring pipelines. Accessible tooling that respects autonomy, privacy, and human final judgment can improve clarity, confidence, and preparation quality without automating hiring outcomes.")

    add_heading(doc, "2. Product capabilities (current build)", 1)
    add_table(
        doc,
        ["Capability", "User value"],
        [
            ("Career profile wizard", "Step-by-step profile (ADHD type, work preferences, support needs, skills/roles) with pause/resume; optional cloud sync"),
            ("Job description simplification", "Paste or upload a posting; AI-assisted plain-language breakdown with ADHD-oriented presentation variants"),
            ("Job fit score (0–100)", "Rule-based matching + ML task-completion estimate with factor breakdown and suggested supports—not a hiring verdict"),
            ("Saved insights", "Compare and revisit scores, day-in-life simulations, and interview prep per role"),
            ("Interview prep", "Guided stages (brain dump, practice questions, feedback) tied to a specific job context"),
            ("Day in the life", "Simulated schedule/expectations to explore role rhythm (clearly labelled as illustrative)"),
            ("Settings & transparency", "Privacy FAQ, data-use explanations, assistive-AI disclaimers"),
        ],
    )
    add_para(doc, "Technical stack (summary): React frontend (Vercel); FastAPI backend (AWS Lambda container + API Gateway); PostgreSQL (AWS RDS) for optional logged-in sessions; Google Gemini API for generative features; Jobs and Skills Australia taxonomy; ADHD research datasets for ML features.")

    add_heading(doc, "3. Resources required", 1)
    add_heading(doc, "People & skills", 2)
    add_bullets(doc, [
        "Product / UX: neuro-inclusive interaction design, consent flows, plain-language copy",
        "Full-stack engineering: React, FastAPI, AWS deployment, API security",
        "Data / ML: feature engineering, model validation, bias review for Job Score",
        "Ethics & compliance: privacy impact, responsible-AI labelling, ACS-aligned governance",
    ])
    add_heading(doc, "Infrastructure & data", 2)
    add_bullets(doc, [
        "Hosting: Vercel (frontend), AWS Lambda/ECR + API Gateway, Render (API fallback)",
        "Data stores: Browser local storage; AWS RDS for cloud-backed profiles/sessions",
        "Third-party services: Google Gemini API (JD simplification, interview prep, day-in-life generation)",
        "Reference data: Australian Jobs & Skills occupation/skill taxonomy; ADHD_4Class and ADHD Task Productivity (Kaggle) datasets for ML layer",
    ])
    add_heading(doc, "Ongoing operational needs", 2)
    add_bullets(doc, [
        "API keys and secrets management (no keys in client builds except optional preview gate)",
        "Monitoring for Gemini/AWS availability, CORS, and model drift",
        "User feedback channel (identified gap in ethics canvas—see section 4)",
    ])

    add_heading(doc, "4. Ethics canvas & ACS Code of Conduct alignment", 1)
    add_para(doc, "Ethics canvas artefact (PG folder): [INSERT LINK TO YOUR MOODLE/PG FOLDER — e.g. NeuroGuide_Ethics_Canvas.pdf]")
    add_para(doc, "Local team copy: NeuroGuide_Ethics_Canvas.pdf (Stellar Minds, 2026).")
    add_para(doc, "The ethics canvas reviews scope, stakeholders, data processing, explainability, human-in-the-loop design, failure modes, and mitigations. The table below maps how NeuroGuide aligns with the ACS Code of Professional Conduct (six core values).")
    add_table(
        doc,
        ["ACS value", "How NeuroGuide addresses it"],
        [
            ("1. Primacy of the public interest", "Product scope is assistive career exploration, not automated hiring or clinical diagnosis. Disclaimers state outputs are not employer-certified. Privacy minimisation and secure storage for sensitive ADHD-related preferences. Users retain decision authority (apply / save / reject)."),
            ("2. Enhancement of quality of life", "Reduces information overload; supports confidence and preparation for underserved neurodivergent IT job seekers. Preserves essential JD facts rather than oversimplifying away material constraints."),
            ("3. Honesty", "Separates original job-description facts from AI interpretation; cautious wording vs absolute claims; labels uncertainty for vague postings; documents dataset and model limitations in the ethics canvas."),
            ("4. Competence", "Uses authoritative Australian occupation data; combines rule-based and ML scoring with documented limitations; secure API design and deployment documentation. Acknowledged gaps: formal feedback/appeal channel and cloud data-deletion workflow still to mature."),
            ("5. Professional development", "Team applies industry-standard stack and responsible-AI review via ethics canvas; handover package supports auditability for assessors and sponsors."),
            ("6. Professionalism", "Respectful UX for neurodivergent users; transparent third-party subprocessors (Gemini, AWS); regulatory awareness (privacy, consent, accessibility, API terms)."),
        ],
    )
    add_para(doc, "Key ethical commitments: data minimisation; consent for profile/cloud sync; explainable job scores; human-in-the-loop for all career decisions; bias testing; no presentation of outputs as clinical or professional career certification.")
    add_para(doc, "Known ethical gaps: limited in-app feedback/objection path; sensitive ADHD profile data requires continued deletion/control tooling; risk of over-trusting AI scores—mitigated through copy, factor breakdowns, and reminders to verify original postings.")

    add_heading(doc, "5. Team — Stellar Minds", 1)
    add_table(
        doc,
        ["Name", "Role (typical)", "Email", "LinkedIn"],
        [
            ("Parit Voruganti", "[e.g. Product / Full-stack lead]", "pvor0002@student.monash.edu", "[INSERT URL]"),
            ("Shibila Thangavelu", "[e.g. Backend / ML]", "stha0086@student.monash.edu", "[INSERT URL]"),
            ("Shreya Srivastava", "[e.g. Frontend / UX]", "ssri0062@student.monash.edu", "[INSERT URL]"),
            ("Chengzhi Hu", "[e.g. Data / QA]", "chuu0022@student.monash.edu", "[INSERT URL]"),
            ("Casper", "[e.g. Engineering]", "shibby776@gmail.com", "[INSERT URL]"),
        ],
    )
    add_para(doc, "Repository: NeuroGuide_TP19 (Monash TP19). Ethics canvas designed by Stellar Minds.")

    add_heading(doc, "6. Demonstration assets", 1)
    add_table(
        doc,
        ["Asset", "Link / access"],
        [
            ("Live build", "https://www.neuroguide.dev (primary domain)"),
            ("Alternate frontend", "https://neuroguide-frontend-iteration-1.vercel.app"),
            ("API (production)", "https://neuroguide-tp19.onrender.com/api/v1"),
            ("Preview password", "[INSERT VITE_SITE_PASSWORD — share via secure handover channel]"),
            ("Product video", "[INSERT YouTube/Vimeo/OneDrive link]"),
            ("Source code", "[INSERT GitHub repo URL or assessor access instructions]"),
        ],
    )
    add_para(doc, "Note: When VITE_SITE_PASSWORD is set, visitors enter the preview gate once per browser session.")

    add_heading(doc, "7. Future sponsors — who and why", 1)
    add_para(doc, "Problem: Early-career IT candidates with ADHD face disproportionate friction reading job ads, comparing roles, and preparing interviews—despite strong technical potential. Sponsorship funds product hardening, user research, and distribution through organisations that already serve this audience.")

    add_heading(doc, "7.1 University & student services", 2)
    add_table(doc, ["Potential sponsor", "Why they would sponsor NeuroGuide"], [
        ("Monash University — Career Connect & Disability Support (DPS)", "Direct alignment with Monash IT graduates: a free, ethics-reviewed tool that helps students before appointments, reducing repeat questions and improving career conversations. Sponsorship could pilot NeuroGuide in one faculty (e.g. IT)."),
        ("Other Australian universities (e.g. RMIT, Deakin, UTS career services)", "Same graduate-employment mission; low integration cost (web app). Sponsorship differentiates their neuro-inclusive employability offer without building in-house AI."),
    ])

    add_heading(doc, "7.2 Industry body & government skills", 2)
    add_table(doc, ["Potential sponsor", "Why they would sponsor NeuroGuide"], [
        ("Australian Computer Society (ACS)", "NeuroGuide maps to the ACS Code of Professional Conduct. Sponsorship positions ACS as supporting responsible assistive AI for members and students—not automated hiring."),
        ("Jobs and Skills Australia", "Product already uses official Australian occupation and skills taxonomy. Sponsorship validates a practical consumer of JSA data for accessibility and early-career pathways."),
    ])

    add_heading(doc, "7.3 Neurodiversity employment & advocacy", 2)
    add_table(doc, ["Potential sponsor", "Why they would sponsor NeuroGuide"], [
        ("Xceptional", "Places neurodivergent talent in tech roles. NeuroGuide prepares candidates upstream—clearer JD understanding and interview prep—so coaches spend less time decoding postings."),
        ("Neurodiversity Hub", "Bridges campuses and inclusive employers. Sponsor-branded instance gives students a consistent pre-employment tool across partner universities."),
        ("Employ for Ability / ADHD Australia", "Mission fit: reduce barriers for neurodivergent Australians entering work. Sponsorship extends reach through trusted community channels with clear assistive (not clinical) disclaimers."),
    ])

    add_heading(doc, "7.4 Inclusive employers (IT & professional services)", 2)
    add_table(doc, ["Potential sponsor", "Why they would sponsor NeuroGuide"], [
        ("Atlassian, Microsoft Australia, Telstra, EY, IBM", "Existing neurodiversity hiring programs. Sponsorship supports employer brand and pipeline quality: candidates understand roles before applying; fewer mismatched applications; not a screening gate."),
        ("GradConnection / early-career platforms", "Large IT graduate audience. Sponsorship adds differentiated accessibility (JD simplification + prep) and supports inclusion reporting on graduate outcomes."),
    ])

    add_heading(doc, "7.5 What sponsorship would fund", 2)
    add_bullets(doc, [
        "API and cloud costs (Gemini, AWS), accessibility (WCAG) audit, and privacy hardening",
        "Formal user feedback and score-challenge workflow (ethics canvas gap)",
        "Pilot studies with a sponsor’s student or candidate cohort (with ethics approval)",
        "Co-branded deployment for one partner (e.g. Monash IT + Career Connect)",
    ])
    add_para(doc, "Why sponsor now? NeuroGuide is a working prototype with national skills data integration, transparent scoring, and documented ethics—ready for a pilot sponsor to move from class project to sustained community tool.")

    add_heading(doc, "8. Document control", 1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ("Version", "1.0 — Handover 2026 S1"),
            ("Team", "Stellar Minds (TP19)"),
            ("Product", "NeuroGuide"),
            ("Ethics canvas", "NeuroGuide_Ethics_Canvas.pdf → [PG folder link]"),
        ],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
