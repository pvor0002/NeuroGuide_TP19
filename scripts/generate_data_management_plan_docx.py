#!/usr/bin/env python3
"""
Generate NeuroGuide Iteration 1 Data Management Plan (.docx).

Requires: python-docx, pillow
Run from repository root:
  python3 scripts/generate_data_management_plan_docx.py
"""

from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Styling aligned with typical Word academic / Stellar Minds report (dark blue headings, Calibri body)
HEADING_BLUE = RGBColor(0x1F, 0x4E, 0x79)  # #1f4e79
BODY_RGB = RGBColor(0x33, 0x33, 0x33)
FONT_NAME = "Calibri"
ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
ASSETS = DATA / "dmp_code_snippets"
OUT_DOCX = DATA / "NeuroGuide_Iteration1_Data_Management_Plan.docx"
FONT_PATH = Path("/System/Library/Fonts/Supplemental/Courier New.ttf")


def shade_cell(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_run_font(
    run,
    size_pt: float,
    bold: bool = False,
    color: RGBColor | None = None,
    italic: bool = False,
) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_heading_custom(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sz = 16 if level == 1 else 14 if level == 2 else 12
    set_run_font(run, sz, bold=True, color=HEADING_BLUE)


def add_para(doc: Document, text: str, bullet: bool = False) -> None:
    style = "List Bullet" if bullet else None
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, 11, color=BODY_RGB)


def add_placeholder(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"[Placeholder: {label}]")
    set_run_font(run, 11, color=RGBColor(0x80, 0x80, 0x80), italic=True)
    p.paragraph_format.space_after = Pt(10)


def render_code_png(out_path: Path, lines: list[str], caption_top: str = "") -> bool:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return False

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pad = 16
    line_h = 18
    title_h = 28 if caption_top else 0
    w = 920
    h = pad * 2 + title_h + len(lines) * line_h + 8

    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype(str(FONT_PATH), 13)
        font_small = ImageFont.truetype(str(FONT_PATH), 11)
    except OSError:
        font = ImageFont.load_default()
        font_small = font

    if caption_top:
        draw.text((pad, pad), caption_top, fill=(31, 78, 121), font=font_small)

    y = pad + title_h
    for line in lines:
        draw.text((pad, y), line, fill=(20, 30, 40), font=font)
        y += line_h

    draw.rectangle([2, 2, w - 3, h - 3], outline=(200, 210, 220))
    img.save(out_path, "PNG")
    return True


def build_code_snippets() -> list[Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []

    snip1 = textwrap.dedent(
        """
        def extract_occupations(workbook) -> list[str]:
            sheet = workbook["Table_1"]
            occupations_to_employment: dict[str, int] = {}
            for row in sheet.iter_rows(min_row=8, values_only=True):
                if not row or len(row) < 2 or row[1] is None:
                    continue
                occupation = normalize_whitespace(str(row[1]))
                ...
                employed_count = int(employed_digits) if employed_digits else 0
                previous_value = occupations_to_employment.get(occupation, 0)
                occupations_to_employment[occupation] = max(previous_value, employed_count)
            ranked = sorted(occupations_to_employment.items(), key=lambda item: (-item[1], item[0]))
            return [occupation for occupation, _ in ranked]
        """
    ).strip().split("\n")

    snip2 = textwrap.dedent(
        """
        def build_skill_labels(action_counter: Counter, minimum_frequency: int = 18) -> list[str]:
            base_labels = {"Communication", "Team Collaboration", "Problem Solving", ...}
            for action_word, frequency in action_counter.items():
                if frequency < minimum_frequency:
                    continue
                label = title_from_action_word(action_word)
                if label:
                    base_labels.add(label)
            return sorted(base_labels)
        """
    ).strip().split("\n")

    snip3 = textwrap.dedent(
        """
        def main() -> None:
            workbook = load_workbook(INPUT_PATH, read_only=True, data_only=True)
            occupations = extract_occupations(workbook)
            action_counter = extract_task_action_words(workbook)
            skill_labels = build_skill_labels(action_counter)
            write_json(SKILLS_OUTPUT_PATH, skill_labels)
            write_json(ROLES_OUTPUT_PATH, occupations)
        """
    ).strip().split("\n")

    specs = [
        (ASSETS / "snippet_extract_occupations.png", snip1, "build_taxonomy_from_excel.py - role list (Table_1)"),
        (ASSETS / "snippet_build_skill_labels.png", snip2, "build_taxonomy_from_excel.py - skill tag synthesis"),
        (ASSETS / "snippet_main_pipeline.png", snip3, "build_taxonomy_from_excel.py - end-to-end pipeline"),
    ]
    for path, lines, cap in specs:
        if render_code_png(path, lines, cap):
            paths.append(path)
    return paths


def add_footer(doc: Document, text: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, 9, color=RGBColor(0x66, 0x66, 0x66))


def add_source_table(doc: Document) -> None:
    headers = [
        "Data source name",
        "Physical access",
        "Source update frequency",
        "System update frequency",
        "Granularity",
        "License and notes",
    ]
    rows = [
        [
            "Jobs and Skills Australia - Occupation profiles workbook "
            "(Occupation profiles data - February 2026.xlsx)",
            "Local file under data/; opened with Python openpyxl (read-only)",
            "Per Jobs and Skills Australia publication cycle (replace file when a new edition is released)",
            "Manual: run python3 scripts/build_taxonomy_from_excel.py after replacing the workbook",
            "Occupation-level rows (Table_1); task statements (Table_3)",
            "Use per Jobs and Skills Australia terms; attribute source in product documentation; "
            "do not imply endorsement.",
        ],
        [
            "Derived: au-role-taxonomy.json",
            "Version-controlled JSON in data/; copied to dist/ on build for static hosting",
            "Regenerated whenever the Excel source is refreshed",
            "Same as workbook refresh; deterministic output for reproducible builds",
            "One string per occupation title (ranked by employment proxy where available)",
            "Derivative of official occupation statistics; document provenance in this DMP.",
        ],
        [
            "Derived: au-skills-taxonomy.json",
            "Version-controlled JSON; build pipeline as above",
            "Regenerated with roles file",
            "Same as workbook refresh",
            "Curated skill chip labels (verbs → categories + transferability tags)",
            "Derived; threshold and mapping logic documented in §3.2.",
        ],
    ]

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        shade_cell(cell, "D9E2F3")
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, 10, bold=True, color=HEADING_BLUE)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    set_run_font(r, 10, color=BODY_RGB)
    doc.add_paragraph()


def main() -> None:
    snippet_paths = build_code_snippets()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_RGB
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Data Management Plan\nNeuroGuide - Iteration 1")
    set_run_font(r, 20, bold=True, color=HEADING_BLUE)
    t.paragraph_format.space_after = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Stellar Minds · FIT5120 TP19")
    set_run_font(sr, 11, color=BODY_RGB)
    doc.add_paragraph()

    add_footer(doc, "NeuroGuide  |  Data Management Plan  |  Stellar Minds TP19")

    # --- 1 ---
    add_heading_custom(doc, "1. Product overview and data role", 1)
    add_para(
        doc,
        "NeuroGuide is a job-support application aimed at reducing cognitive load for neurodivergent job seekers, "
        "including structured capture of skills and background, clearer presentation of job information, and "
        "(in later iterations) suitability review and interview preparation. This iteration focuses on "
        "deterministic, auditable reference data for tagging and prototype UI flows.",
    )
    add_para(
        doc,
        "Iteration 1 scope includes: guided skill and background capture, accessibility-oriented input patterns "
        "where implemented in the prototype, and a job-description simplification experience - delivered primarily "
        "as front-end flows backed by curated static taxonomies rather than live model APIs.",
    )
    add_para(doc, "Functions supported by data in this iteration:", bullet=False)

    ft = doc.add_table(rows=3, cols=3)
    ft.rows[0].cells[0].text = "Functional module"
    ft.rows[0].cells[1].text = "Main input data"
    ft.rows[0].cells[2].text = "Output / use"
    for c in ft.rows[0].cells:
        shade_cell(c, "D9E2F3")
        for p in c.paragraphs:
            for run in p.runs:
                set_run_font(run, 10, bold=True, color=HEADING_BLUE)
    data_rows = [
        (
            "Guided skill & background profile",
            "User input + au-skills-taxonomy.json + au-role-taxonomy.json",
            "Searchable tags/chips; local draft persisted in browser storage",
        ),
        (
            "Job description simplification UI (prototype)",
            "User-pasted or attached text; optional future NLP API (Iteration 2+)",
            "Structured sections in-app; document export may be added in a later iteration",
        ),
    ]
    for i, (a, b, c) in enumerate(data_rows, start=1):
        ft.rows[i].cells[0].text = a
        ft.rows[i].cells[1].text = b
        ft.rows[i].cells[2].text = c
        for cell in ft.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, 10, color=BODY_RGB)
    doc.add_paragraph()

    add_para(
        doc,
        "Data categories in Iteration 1: (1) Official Australian occupation and task statistics "
        "(Excel workbook). (2) Derived static taxonomies (JSON) for UI tag libraries. "
        "(3) Ephemeral user-entered profile fields (browser only in the current prototype). "
        "Later iterations may add further datasets, suitability scoring (e.g. vector similarity over skill embeddings), "
        "and external NLP or speech services subject to privacy review.",
        bullet=False,
    )
    add_placeholder(doc, "Screenshot - product landing or iteration scope overview")

    # --- 2 ---
    add_heading_custom(doc, "2. Data sources and update strategy", 1)
    add_heading_custom(doc, "2.1 Open data and project source table", 2)
    add_source_table(doc)
    add_placeholder(doc, "Screenshot - Jobs and Skills Australia download or attribution page (when final URL confirmed)")

    add_heading_custom(doc, "2.2 Principles for selecting data sources", 2)
    for line in [
        "Prefer open, documented sources with clear licensing and Australian labour-market relevance for occupation and skills alignment.",
        "Separate official reference statistics (workbook) from derived artefacts (JSON) so updates are reproducible via a single script.",
        "Defer high-risk processing (résumé storage, cloud speech, external LLMs) to later iterations with explicit privacy review.",
        "Keep Iteration 1 deterministic: the taxonomy build is scriptable and suitable for code review and regression checks.",
    ]:
        add_para(doc, line, bullet=True)

    # --- 3 ---
    add_heading_custom(doc, "3. Data flow, data cleansing and transformation", 1)
    add_heading_custom(doc, "3.1 High-level data flow", 2)
    for line in [
        "Analyst or developer obtains the latest Jobs and Skills Australia occupation profiles Excel file and places it in data/.",
        "Python loads Table_1 (occupation titles and employment-related fields) and Table_3 (task statements) using openpyxl in read-only mode.",
        "Cleansing and feature engineering produce two sorted string lists, written as UTF-8 JSON with stable indentation.",
        "The front-end build copies public assets; the React profile flow fetches au-role-taxonomy.json and au-skills-taxonomy.json at runtime.",
        "User selections are combined with free-text fields; draft state is held client-side (localStorage) in the prototype.",
    ]:
        add_para(doc, line, bullet=True)

    add_heading_custom(doc, "3.2 Wrangling and cleansing (Python / data science operations)", 2)

    add_heading_custom(doc, "3.2.1 Ingestion and parsing", 3)
    add_para(
        doc,
        "The workbook is treated as semi-structured: logical tables begin at fixed row offsets (e.g. min_row=8). "
        "Rows with missing occupation or task fields are skipped. Text fields are normalised with regular expressions "
        "to collapse internal whitespace, reducing duplicate keys caused only by formatting differences.",
    )

    add_heading_custom(doc, "3.2.2 Frequency analysis and thresholding (skills)", 3)
    add_para(
        doc,
        "Task statements are scanned from Table_3. The first alphabetic token of each task line is treated as a "
        "leading action verb (proxy for task type). Frequencies are aggregated with collections.Counter - a standard "
        "exploratory step to identify dominant task verbs in the corpus. Only verbs meeting a minimum count threshold "
        "(default 18) are promoted to skill labels, which suppresses long-tail noise while retaining statistically "
        "supported constructs.",
    )

    add_heading_custom(doc, "3.2.3 Controlled vocabulary mapping and deduplication (roles)", 3)
    add_para(
        doc,
        "Occupation titles are deduplicated in dictionary space; when duplicate labels appear with differing employment "
        "count strings, the maximum parsed numeric employment proxy is retained so ranking favours the stronger record. "
        "The final role list is sorted by descending employment proxy then alphabetically - a transparent, auditable rule.",
    )
    add_para(
        doc,
        "Verb-to-skill-label mapping uses a curated dictionary (domain-informed categorical encoding). Unmapped verbs "
        "contribute no label, which is a deliberate bias–variance trade-off: the UI stays short and interpretable.",
    )

    add_heading_custom(doc, "3.2.4 Client-side consumption", 3)
    add_para(
        doc,
        "The web client loads JSON arrays and filters options against user search queries. This is lightweight "
        "wrangling in the presentation layer (substring match, caps on visible results) to protect responsiveness.",
    )

    add_heading_custom(doc, "3.2.5 Techniques used (summary)", 3)
    tech = doc.add_table(rows=7, cols=2)
    tech.rows[0].cells[0].text = "Technique"
    tech.rows[0].cells[1].text = "Application in NeuroGuide Iteration 1"
    shade_cell(tech.rows[0].cells[0], "E7E6E6")
    shade_cell(tech.rows[0].cells[1], "E7E6E6")
    pairs = [
        ("String cleaning", "regex whitespace normalisation; numeric token extraction from employment fields"),
        ("Frequency filtering", "Counter on leading task verbs; minimum_frequency cutoff"),
        ("Categorical mapping", "verb → human-readable skill chip via dictionary lookup"),
        ("Deduplication / conflict resolution", "max employment count per occupation string"),
        ("Serialisation", "json.dumps with ensure_ascii=False, stable key order via sorted()"),
        ("Reproducibility", "single entry-point script; pinned input path; print counts on success"),
    ]
    for i, (a, b) in enumerate(pairs, start=1):
        tech.rows[i].cells[0].text = a
        tech.rows[i].cells[1].text = b
        for c in tech.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    set_run_font(r, 10, color=BODY_RGB)
    doc.add_paragraph()

    add_heading_custom(doc, "Code excerpts (Python)", 2)
    add_para(doc, "Key excerpts from scripts/build_taxonomy_from_excel.py:", bullet=False)
    for pth in snippet_paths:
        doc.add_picture(str(pth), width=Inches(6.2))
        doc.add_paragraph()
    if not snippet_paths:
        add_placeholder(doc, "Screenshot - scripts/build_taxonomy_from_excel.py in IDE (if image generation unavailable)")

    add_heading_custom(doc, "3.3 Transposition into the application", 2)
    add_para(
        doc,
        "Clean JSON arrays transpose directly into UI affordances: searchable multi-select tags for roles and skills. "
        "The pattern matches standard practice: cleaned, stable artefacts exposed to the client in a shape the UI "
        "can consume directly - here via static hosting and fetch() rather than a dynamic query API.",
    )

    add_heading_custom(doc, "3.4 Storage and archiving", 2)
    for line in [
        "Source Excel and derived JSON live under data/ and are versioned in Git.",
        "Production build copies JSON into dist/ for static deployment; tag outputs with commit hash in release notes.",
        "Retain superseded Excel files with a dated filename or branch tag before overwriting.",
    ]:
        add_para(doc, line, bullet=True)
    add_placeholder(doc, "Iteration 2 - S3 / database ERD or bucket policy screenshot for résumé and profile API")

    # --- 4 ---
    add_heading_custom(doc, "4. Data update frequency and iteration process", 1)
    for line in [
        "Static taxonomies: update when Jobs and Skills Australia publishes a new occupation profiles workbook; re-run the build script and QA the chip counts in the UI.",
        "User draft profiles: no central persistence in Iteration 1; clearing browser data clears drafts.",
        "Teaching feedback: incorporate Monash / client feedback into threshold or mapping tweaks and document changes in this DMP.",
    ]:
        add_para(doc, line, bullet=True)
    add_placeholder(doc, "Iteration 2 - planned cadence for ADHD modelling datasets and API monitoring")

    # --- 5 ---
    add_heading_custom(doc, "5. Storage strategy", 1)
    st = doc.add_table(rows=6, cols=4)
    st.rows[0].cells[0].text = "Data type"
    st.rows[0].cells[1].text = "Persisted?"
    st.rows[0].cells[2].text = "Location"
    st.rows[0].cells[3].text = "Archiving"
    for c in st.rows[0].cells:
        shade_cell(c, "D9E2F3")
        for p in c.paragraphs:
            for r in p.runs:
                set_run_font(r, 10, bold=True, color=HEADING_BLUE)
    storage_rows = [
        ("Raw occupation workbook", "Yes", "data/*.xlsx", "Git + dated copies"),
        ("au-role-taxonomy.json / au-skills-taxonomy.json", "Yes", "data/ → dist/ on build", "Git diff + changelog entry"),
        ("User profile draft (prototype)", "Optional local", "Browser localStorage", "User-controlled; not server-backed"),
        ("Internal product / requirements notes", "As applicable", "Team repository or wiki", "Version with release or submission package"),
        ("Iteration 2 - embeddings / models", "[Planned]", "[TBD]", "[Placeholder: model registry / versioned artifacts]"),
    ]
    for i, row in enumerate(storage_rows, start=1):
        for j, val in enumerate(row):
            st.rows[i].cells[j].text = val
            for p in st.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    set_run_font(r, 10, color=BODY_RGB)
    doc.add_paragraph()

    # --- 6 ---
    add_heading_custom(doc, "6. Data modelling and analysis", 1)
    add_heading_custom(doc, "6.1 Conceptual and logical view", 2)
    add_para(doc, "Core entities for Iteration 1:", bullet=False)
    mt = doc.add_table(rows=6, cols=2)
    mt.rows[0].cells[0].text = "Entity"
    mt.rows[0].cells[1].text = "Attributes (illustrative)"
    shade_cell(mt.rows[0].cells[0], "E7E6E6")
    shade_cell(mt.rows[0].cells[1], "E7E6E6")
    ents = [
        ("OccupationTag", "title string; optional employment rank from ETL"),
        ("SkillTag", "label string; sourced from mapped high-frequency verbs + fixed transferability set"),
        ("UserProfileDraft", "selected roles, skills, rhythms, support prefs; ephemeral in prototype"),
        ("JobPostingText", "raw paste / extract; future: sections from NLP pipeline"),
        ("DataSourceRegistry", "documented in this DMP; not a DB table in Iteration 1"),
    ]
    for i, (a, b) in enumerate(ents, start=1):
        mt.rows[i].cells[0].text = a
        mt.rows[i].cells[1].text = b
        for c in mt.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    set_run_font(r, 10, color=BODY_RGB)
    doc.add_paragraph()
    add_para(
        doc,
        "Later iterations may introduce user skill vectors, job requirement vectors, similarity-based suitability "
        "scores, and optional behavioural or attention-profile inputs to tailor content presentation.",
        bullet=False,
    )
    add_placeholder(doc, "Entity-relationship diagram or vector model diagram - Iteration 2")

    add_heading_custom(doc, "6.2 Modelling techniques and rationale", 2)
    add_para(
        doc,
        "Iteration 1 uses descriptive analytics only: frequency distributions of task-leading verbs and deterministic "
        "rules for label synthesis. This provides an interpretable baseline before predictive models consume "
        "the same tags.",
    )
    add_placeholder(doc, "Iteration 2 - scikit-learn cosine similarity / feature pipeline description")

    # --- 7 ---
    add_heading_custom(doc, "7. Ethical, legal and privacy issues", 1)
    add_heading_custom(doc, "7.1 Data ethics and AI considerations", 2)
    for line in [
        "Occupation statistics describe populations, not individuals; still, recommendations built on them must avoid implying certainty about any one user.",
        "Verb→skill mapping encodes design choices that may under-represent niche tasks; monitor inclusivity feedback.",
        "Future NLP simplification must avoid over-confident rewriting of legal duties or safety-critical requirements.",
    ]:
        add_para(doc, line, bullet=True)

    add_heading_custom(doc, "7.2 Legal and licensing", 2)
    add_para(
        doc,
        "Comply with Jobs and Skills Australia publication terms; maintain attribution in-app or documentation. "
        "Any future third-party APIs (for example cloud language or speech services) must be onboarded only under "
        "compliant terms, keys, and privacy review. Iteration 1 limits scope by using static taxonomies for tag pickers.",
    )

    add_heading_custom(doc, "7.3 Privacy", 2)
    add_para(
        doc,
        "The current prototype does not send profile payloads to a team-controlled backend. Iteration 2+ plans "
        "(résumé uploads, voice, suitability persistence) require encryption in transit, access controls, and clear "
        "retention policy - placeholder for security architecture screenshot and DPIA summary.",
    )
    add_placeholder(doc, "Screenshot - privacy posture / threat model workshop whiteboard (Iteration 2)")

    # --- 8 ---
    add_heading_custom(doc, "8. Alignment with problem domain and user needs", 1)
    add_para(
        doc,
        "Structured occupation and skill tags support users who benefit from constrained choices instead of open-ended "
        "essays. Keeping job text processing transparent and incremental reduces reliance on heavy cognitive "
        "organisation before value is delivered. Data design priorities are clarity, reproducibility, and minimal "
        "collection of sensitive information in this iteration.",
    )
    add_placeholder(doc, "Screenshot - user journey or needs summary (optional)")

    # --- 9 ---
    add_heading_custom(doc, "9. Summary", 1)
    sm = doc.add_table(rows=6, cols=2)
    sm.rows[0].cells[0].text = "Area"
    sm.rows[0].cells[1].text = "Summary"
    for c in sm.rows[0].cells:
        shade_cell(c, "D9E2F3")
        for p in c.paragraphs:
            for r in p.runs:
                set_run_font(r, 10, bold=True, color=HEADING_BLUE)
    summaries = [
        ("Data sources", "Jobs and Skills Australia occupation workbook; derived JSON; repository artefacts and this DMP."),
        ("Processes", "Python ETL with openpyxl; regex cleaning; frequency thresholds; JSON export; static hosting."),
        ("Modelling", "Descriptive frequency and rule-based mapping; predictive suitability deferred to Iteration 2."),
        ("Governance", "This DMP; Git history; placeholders for DPIA and API terms as scope expands."),
        ("Live document", "Update each iteration as sources, models, and compliance posture change."),
    ]
    for i, (a, b) in enumerate(summaries, start=1):
        sm.rows[i].cells[0].text = a
        sm.rows[i].cells[1].text = b
        for c in sm.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    set_run_font(r, 10, color=BODY_RGB)

    doc.add_paragraph()
    add_para(
        doc,
        "Primary data artefacts in this iteration: Occupation profiles data - February 2026.xlsx (source workbook), "
        "au-role-taxonomy.json, and au-skills-taxonomy.json (derived taxonomies).",
        bullet=False,
    )

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
